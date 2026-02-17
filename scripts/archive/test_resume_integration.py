
import requests
import docx
import os
import sys

# Add project root to path if needed (not needed for requests)

def create_resume():
    try:
        doc = docx.Document()
        doc.add_heading('John Doe', 0)
        doc.add_paragraph('Senior Python Developer | 5+ Years Experience')
        doc.add_paragraph('San Francisco, CA | john@example.com')
        doc.add_heading('Skills', level=1)
        doc.add_paragraph('Python, Django, FastAPI, PostgreSQL, React, AWS, Docker, Kubernetes')
        doc.add_heading('Experience', level=1)
        doc.add_paragraph('Senior Backend Engineer at TechCorp (2020-Present)')
        doc.add_paragraph('Built scalable APIs using FastAPI and Python. Implemented resume parsing.')
        doc.save('test_resume.docx')
        print("✅ Created test_resume.docx")
    except Exception as e:
        print(f"❌ Failed to create DOCX: {e}")
        sys.exit(1)

def test_upload():
    # Inside container, API is at localhost:8000
    url = "http://localhost:8000/api/match-resume" 
    
    if not os.path.exists('test_resume.docx'):
        print("❌ Resume file missing!")
        return

    try:
        files = {'file': ('test_resume.docx', open('test_resume.docx', 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        print(f"🚀 Uploading to {url}...")
        response = requests.post(url, files=files)
        
        if response.status_code == 200:
            data = response.json()
            matches = data.get('matches', [])
            print(f"🎉 Success! Found {len(matches)} matches.")
            
            if matches:
                top = matches[0]
                print(f"\n🏆 TOP MATCH: {top['job']['title']} at {top['job']['company']}")
                print(f"   Total Score: {top['score']['total_score']}%")
                print(f"   Breakdown: {top['score']['breakdown']}")
                
                # Check 3rd match to see variation
                if len(matches) > 2:
                    third = matches[2]
                    print(f"\n🥉 3RD MATCH: {third['job']['title']}")
                    print(f"   Total Score: {third['score']['total_score']}%")
            else:
                print("⚠️ No matches found. Check database population.")
                
        else:
            print(f"❌ API Error: {response.status_code} - {response.text}")
            
    except Exception as e:
        print(f"❌ Request Failed: {e}")

if __name__ == "__main__":
    create_resume()
    test_upload()
